"""HTML/PDF/DOCX export and social scorecard generation for Career Intelligence Reports."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path

from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.schemas.career_intelligence_report import CareerIntelligenceReportResponse
from app.schemas.report_export import ReportScorecardResponse

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates" / "export"

_jinja = Environment(
    loader=FileSystemLoader(str(TEMPLATE_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)


def _format_date(value: datetime | None) -> str:
    if not value:
        return datetime.utcnow().strftime("%B %d, %Y")
    return value.strftime("%B %d, %Y")


def _portfolio_lines(mix: dict[str, int]) -> list[str]:
    total = sum(mix.values()) or 1
    lines = []
    for key in ("BUILD", "BLEND", "BOT"):
        count = mix.get(key, 0)
        if count:
            pct = round((count / total) * 100)
            lines.append(f"{pct}% {key}")
    return lines


def _kpi_map(report: CareerIntelligenceReportResponse) -> dict[str, str]:
    return {kpi.label: kpi.value for kpi in report.overview.kpis}


def _snapshot_map(report: CareerIntelligenceReportResponse) -> dict[str, str]:
    return {item.label: item.value for item in report.overview.career_snapshot}


def _hours_automatable_per_week(report: CareerIntelligenceReportResponse) -> float:
    daily_by_name = {task.name: task.hours_per_week for task in report.daily_work.tasks}
    freed = 0.0
    for analysis in report.task_routing.analyses:
        if (analysis.category or "").upper() != "BOT":
            continue
        hours = daily_by_name.get(analysis.task_title or "", 0.0)
        potential = float(analysis.auto_potential or 50) / 100
        freed += hours * potential * 0.65
    return round(freed, 1)


def build_pdf_export_context(
    report: CareerIntelligenceReportResponse,
    *,
    recipient_name: str,
    job_title: str | None = None,
) -> dict:
    """Map API report model to the flat template contract used by report_pdf.html."""
    readiness = report.ai_readiness
    kpis = _kpi_map(report)
    snapshot = _snapshot_map(report)
    mix = readiness.portfolio_mix
    role = job_title or snapshot.get("Current Role") or report.before_after.role_today

    automation_raw = kpis.get("Automation %", "—")
    automation_pct = automation_raw.replace("%", "") if automation_raw != "—" else "—"
    experience_raw = snapshot.get("Experience", "—")
    experience_years = experience_raw.replace(" years", "") if experience_raw != "—" else "—"

    overview = {
        "job_title": role,
        "industry": snapshot.get("Industry", "—"),
        "overall_score": readiness.overall_score,
        "tasks_analyzed": kpis.get("Tasks Analyzed", str(len(report.task_routing.analyses))),
        "competency_count": kpis.get("Competencies", str(len(report.competencies))),
        "automation_pct": automation_pct,
        "career_risk": readiness.career_risk,
        "ai_tools_count": str(len(report.ai_toolkit)),
        "experience_years": experience_years,
        "profession_summary": report.overview.insight,
    }

    ai_readiness = {
        "overall_score": readiness.overall_score,
        "tier_label": readiness.tier_label,
        "tier_description": readiness.summary,
        "dimensions": [{"name": dim.subject, "score": dim.score} for dim in readiness.dimensions],
        "strengths": [item.title for item in readiness.strengths],
        "improvement_areas": [item.title for item in readiness.improvements],
    }

    career_identity = {
        "identity_title": report.career_identity.title,
        "confidence_pct": readiness.overall_score,
        "executive_summary": report.career_identity.narrative,
        "growth_strategy": report.career_identity.closing_note,
    }

    task_routing = [
        {
            "task_title": analysis.task_title,
            "category": analysis.category,
            "rationale": analysis.rationale,
            "reason": analysis.reason,
            "next_actions": list(analysis.next_actions or []),
            "recommended_tools": list(analysis.recommended_tools or []),
        }
        for analysis in report.task_routing.analyses
    ]

    competencies = [
        {
            "category": group.title,
            "items": [
                {
                    "name": item.name,
                    "importance": item.importance,
                    "expected_level": item.growth,
                }
                for item in group.items
            ],
        }
        for group in report.competencies
    ]

    daily_work = {
        "total_hours_per_week": report.daily_work.total_hours,
        "tasks": [
            {
                "title": task.name,
                "hours_per_week": task.hours_per_week,
                "complexity": task.criticality or "—",
                "ai_assistance": task.ai_usage,
            }
            for task in report.daily_work.tasks
        ],
    }

    learning_roadmap = [
        {
            "horizon": phase.period,
            "title": phase.period,
            "items": [item.title for item in phase.items],
        }
        for phase in report.upskill_roadmap
    ]

    action_plan = {
        "start": [item.text for item in report.action_plan.start_doing],
        "automate": [item.text for item in report.action_plan.automate_with_ai],
        "learn": [item.text for item in report.action_plan.learn_next],
    }

    ai_toolkit = [tool.model_dump(mode="json") for tool in report.ai_toolkit]
    cost_roi = report.cost_roi.model_dump(mode="json")
    cost_roi["hours_automatable_per_week"] = _hours_automatable_per_week(report)

    return {
        "recipient_name": recipient_name,
        "generated_date": _format_date(report.generated_at),
        "report_version": report.report_version,
        "assessment_id": str(report.assessment_id),
        "strategic_note": report.strategic_note,
        "overview": overview,
        "ai_readiness": ai_readiness,
        "career_identity": career_identity,
        "task_routing": task_routing,
        "competencies": competencies,
        "daily_work": daily_work,
        "learning_roadmap": learning_roadmap,
        "action_plan": action_plan,
        "ai_toolkit": ai_toolkit,
        "cost_roi": cost_roi,
        "build_count": mix.get("BUILD", 0),
        "blend_count": mix.get("BLEND", 0),
        "bot_count": mix.get("BOT", 0),
    }


def generate_scorecard(
    report: CareerIntelligenceReportResponse,
    *,
    job_title: str | None = None,
) -> ReportScorecardResponse:
    readiness = report.ai_readiness
    role = job_title or report.career_identity.title.replace("AI-Augmented ", "")
    mix = readiness.portfolio_mix
    total = sum(mix.values()) or 1
    bot_pct = round((mix.get("BOT", 0) / total) * 100)
    blend_pct = round((mix.get("BLEND", 0) / total) * 100)
    build_pct = round((mix.get("BUILD", 0) / total) * 100)

    headline = (
        f"AI Readiness {readiness.overall_score}/100 — {readiness.tier_label} "
        f"({role})"
    )
    hashtags = ["AIReadiness", "CareerShift", "FutureOfWork", "AICareer"]

    linkedin_text = (
        f"I just completed my CareerShift AI Career Intelligence Assessment.\n\n"
        f"Score: {readiness.overall_score}/100 ({readiness.tier_label} tier)\n"
        f"Role focus: {role}\n"
        f"3B portfolio mix — BUILD {build_pct}% · BLEND {blend_pct}% · BOT {bot_pct}%\n"
        f"Career risk: {readiness.career_risk} · Opportunity: {readiness.career_opportunity}\n\n"
        f"{report.strategic_note}\n\n"
        f"{' '.join('#' + tag for tag in hashtags)}"
    )

    twitter_text = (
        f"CareerShift AI Readiness: {readiness.overall_score}/100 ({readiness.tier_label}). "
        f"{role} — {bot_pct}% tasks BOT-ready, {blend_pct}% BLEND. "
        f"Building an AI-augmented career. {' '.join('#' + tag for tag in hashtags[:3])}"
    )
    if len(twitter_text) > 280:
        twitter_text = (
            f"AI Readiness {readiness.overall_score}/100 ({readiness.tier_label}). "
            f"{bot_pct}% BOT · {blend_pct}% BLEND · {build_pct}% BUILD. "
            f"#CareerShift #AIReadiness"
        )

    return ReportScorecardResponse(
        headline=headline,
        linkedin_text=linkedin_text.strip(),
        twitter_text=twitter_text.strip(),
        hashtags=hashtags,
    )


def render_report_html(
    report: CareerIntelligenceReportResponse,
    *,
    recipient_name: str | None = None,
    job_title: str | None = None,
) -> str:
    template = _jinja.get_template("report_pdf.html")
    context = build_pdf_export_context(
        report,
        recipient_name=recipient_name or "Professional",
        job_title=job_title,
    )
    return template.render(**context)


def render_toolkit_html(
    report: CareerIntelligenceReportResponse,
    *,
    job_title: str | None = None,
) -> str:
    template = _jinja.get_template("ai_toolkit.html")
    return template.render(
        report=report,
        job_title=job_title or report.before_after.role_today,
        generated_date=_format_date(report.generated_at),
    )


def html_to_pdf(html: str) -> bytes:
    from xhtml2pdf import pisa

    buffer = BytesIO()
    status = pisa.CreatePDF(html, dest=buffer, encoding="utf-8")
    if status.err:
        raise RuntimeError("PDF generation failed")
    return buffer.getvalue()


def render_report_pdf(
    report: CareerIntelligenceReportResponse,
    *,
    recipient_name: str | None = None,
    job_title: str | None = None,
) -> bytes:
    html = render_report_html(
        report,
        recipient_name=recipient_name,
        job_title=job_title,
    )
    return html_to_pdf(html)


def _add_bullet_list(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def render_report_docx(
    report: CareerIntelligenceReportResponse,
    *,
    recipient_name: str | None = None,
    job_title: str | None = None,
) -> bytes:
    from docx import Document

    context = build_pdf_export_context(
        report,
        recipient_name=recipient_name or "Professional",
        job_title=job_title,
    )
    overview = context["overview"]
    readiness = context["ai_readiness"]
    identity = context["career_identity"]

    doc = Document()
    doc.add_heading("Career Intelligence Report", 0)
    doc.add_paragraph(
        f"Prepared for {context['recipient_name']} · {overview['job_title']} · {overview['industry']}\n"
        f"Generated {context['generated_date']} · Version {context['report_version']} · "
        f"Assessment {context['assessment_id'][:8]}…"
    )

    doc.add_heading("AI Readiness Score", level=1)
    doc.add_paragraph(
        f"{readiness['overall_score']}/100 — {readiness['tier_label']} — {readiness['tier_description']}"
    )

    doc.add_heading("Executive Overview", level=1)
    if context["strategic_note"]:
        doc.add_paragraph(f"Strategic Note: {context['strategic_note']}")
    if identity["executive_summary"]:
        doc.add_paragraph(identity["executive_summary"])

    doc.add_paragraph(
        f"Overall Score: {overview['overall_score']} | Tasks: {overview['tasks_analyzed']} | "
        f"Competencies: {overview['competency_count']} | Automation: {overview['automation_pct']}% | "
        f"Career Risk: {overview['career_risk']} | AI Tools: {overview['ai_tools_count']}"
    )
    doc.add_paragraph(
        f"Experience: {overview['experience_years']} years · "
        f"3B mix: {context['build_count']} BUILD · {context['blend_count']} BLEND · {context['bot_count']} BOT"
    )
    if overview.get("profession_summary"):
        doc.add_paragraph(f"Profession Summary: {overview['profession_summary']}")

    doc.add_heading("AI Readiness Dimensions", level=1)
    for dim in readiness["dimensions"]:
        doc.add_paragraph(f"{dim['name']}: {dim['score']}/100")

    if readiness["strengths"]:
        doc.add_heading("Strengths", level=2)
        _add_bullet_list(doc, readiness["strengths"])
    if readiness["improvement_areas"]:
        doc.add_heading("Improvement Areas", level=2)
        _add_bullet_list(doc, readiness["improvement_areas"])

    doc.add_heading("Career Identity", level=1)
    doc.add_paragraph(
        f"{identity['identity_title']} ({identity['confidence_pct']}% alignment)"
    )
    if identity["growth_strategy"]:
        doc.add_paragraph(identity["growth_strategy"])

    doc.add_heading("3B Task Analysis", level=1)
    for task in context["task_routing"]:
        doc.add_heading(task["task_title"], level=2)
        doc.add_paragraph(f"Category: {task['category']}")
        if task.get("rationale"):
            doc.add_paragraph(task["rationale"])
        if task.get("reason"):
            doc.add_paragraph(task["reason"])
        if task.get("next_actions"):
            _add_bullet_list(doc, task["next_actions"])
        if task.get("recommended_tools"):
            doc.add_paragraph(f"Tools: {', '.join(task['recommended_tools'])}")

    doc.add_heading("Competencies", level=1)
    for group in context["competencies"]:
        doc.add_heading(group["category"], level=2)
        for item in group["items"]:
            line = item["name"]
            if item.get("importance"):
                line += f" — {item['importance']}"
            if item.get("expected_level"):
                line += f" ({item['expected_level']})"
            doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Daily Work Profile", level=1)
    doc.add_paragraph(f"Total hours/week analyzed: {context['daily_work']['total_hours_per_week']}")
    for task in context["daily_work"]["tasks"]:
        doc.add_paragraph(
            f"{task['title']}: {task['hours_per_week']}h/week · "
            f"Complexity: {task['complexity']} · AI usage: {task.get('ai_assistance') or '—'}"
        )

    doc.add_heading("Learning Roadmap", level=1)
    for phase in context["learning_roadmap"]:
        doc.add_heading(f"{phase['horizon']} — {phase['title']}", level=2)
        _add_bullet_list(doc, phase["items"])

    doc.add_heading("Action Plan", level=1)
    action_plan = context["action_plan"]
    if action_plan.get("start"):
        doc.add_heading("Start", level=2)
        _add_bullet_list(doc, action_plan["start"])
    if action_plan.get("automate"):
        doc.add_heading("Automate", level=2)
        _add_bullet_list(doc, action_plan["automate"])
    if action_plan.get("learn"):
        doc.add_heading("Learn", level=2)
        _add_bullet_list(doc, action_plan["learn"])

    doc.add_heading("Recommended AI Toolkit", level=1)
    for tool in context["ai_toolkit"]:
        rank = tool.get("priority_rank")
        label = tool.get("priority_label") or ""
        rank_text = f"#{rank} {label}".strip() if rank else "—"
        why_parts = []
        if tool.get("priority_reason"):
            why_parts.append(f"Priority: {tool['priority_reason']}")
        if tool.get("use_case"):
            why_parts.append(f"Why: {tool['use_case']}")
        doc.add_paragraph(
            f"{rank_text} · {tool['name']} · {tool.get('category', 'Recommended')}\n"
            + " ".join(why_parts)
        )

    hours_auto = context["cost_roi"].get("hours_automatable_per_week")
    if hours_auto is not None:
        doc.add_heading("Automation Opportunity", level=1)
        doc.add_paragraph(f"Estimated automatable hours/week (BOT tasks): {hours_auto}")

    doc.add_paragraph(
        f"CareerShift Career Intelligence Report · Confidential · Generated from your assessment data only.\n"
        f"© CareerShift · {context['generated_date']}"
    )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
